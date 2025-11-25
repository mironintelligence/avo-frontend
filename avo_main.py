from fastapi import FastAPI, Request, HTTPException
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from datetime import datetime
import io, os, time
from typing import Dict, Any, Literal

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from openai import OpenAI

# ------------------ AVO AI: LIMIT SYSTEM ------------------ #
PETITION_LIMIT = 5             # günlük 5 dilekçe
CHAT_LIMIT_SECONDS = 600       # günlük 10 dk sohbet

usage = {}  # ip -> {date, petitions, chat_seconds}

def today():
    return time.strftime("%Y-%m-%d")

def get_usage(ip):
    if ip not in usage or usage[ip]["date"] != today():
        usage[ip] = {"date": today(), "petitions": 0, "chat_seconds": 0}
    return usage[ip]

# ------------------ OPENAI ------------------ #
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# ------------------ LOAD LIBRA AI CATALOG ------------------ #
# KULLANDIĞIN KODU BİREBİR ALIYORUM, AMA AVO FORMATINA UYARLIYORUM
from catalog_data import CATALOG  # <- bunu ayıracağız

# ---------------- FASTAPI ---------------- #
app = FastAPI(title="AVO AI Legal Writer")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], allow_methods=["*"], allow_headers=["*"]
)

# ---------------- MODELS ---------------- #
class PreviewRequest(BaseModel):
    template_key: str
    values: Dict[str, Any]
    language: Literal["TR", "EN"] = "TR"

class ExportRequest(PreviewRequest):
    format: Literal["docx"] = "docx"

# ---------------- HELPERS ---------------- #
def get_template(key):
    for cat in CATALOG.values():
        for t in cat:
            if t["key"] == key:
                return t
    return None

def build_prompt(tpl, req):
    lang = "Turkish" if req.language == "TR" else "English"

    SYSTEM = """
You are AVO AI Legal Writer.
Generate a formal Turkish legal petition.

IMPORTANT:
- Always add this sentence at the end:
"Bu bir hukuki tavsiye değildir."
- Keep content structured.
- Respect typical Turkish dilekçe format.
"""

    def items(d):
        s = []
        for k, v in d.items():
            if v: s.append(f"- {k}: {v}")
        return "\n".join(s)

    prompt = f"""
{SYSTEM}

LANGUAGE: {lang}
TEMPLATE: {tpl['title']}
CASE TYPE: {tpl['case_type']}
POLICY: {tpl['policy']}

USER INPUTS:
{items(req.values)}

Output sections:
#HEADER
#SUMMARY
#LEGAL_BASIS
#RESULT_REQUESTS
#ATTACHMENTS
"""

    return prompt

def parse_ai(text: str):
    keys = ["#HEADER", "#SUMMARY", "#LEGAL_BASIS", "#RESULT_REQUESTS", "#ATTACHMENTS"]
    out = {k.lower().replace("#",""): "" for k in keys}

    for i in range(len(keys)):
        start_key = keys[i]
        start = text.find(start_key)
        if start == -1:
            continue

        end = text.find(keys[i+1]) if i+1 < len(keys) else len(text)

        section = text[start+len(start_key):end].strip()
        out[start_key.lower().replace("#","")] = section

    return out

# ---------------- ROUTES ---------------- #

@app.get("/catalog")
def list_catalog():
    out = []
    for cat, items in CATALOG.items():
        out.append({
            "category": cat,
            "items": [
                {"key": t["key"], "title": t["title"], "case_type": t["case_type"]}
                for t in items
            ]
        })
    return out


@app.get("/fields/{template_key}")
def get_fields(template_key: str):
    tpl = get_template(template_key)
    if not tpl:
        raise HTTPException(404, "Şablon yok")
    return {"title": tpl["title"], "fields": tpl["fields"]}


@app.post("/preview")
def preview(req: PreviewRequest, request: Request):
    ip = request.client.host
    u = get_usage(ip)

    # limit kontrol
    if u["petitions"] >= PETITION_LIMIT:
        raise HTTPException(429, "Günlük 5 dilekçe limitine ulaştınız.")

    u["petitions"] += 1

    tpl = get_template(req.template_key)
    if not tpl:
        raise HTTPException(400, "Geçersiz template_key")

    prompt = build_prompt(tpl, req)

    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        temperature=0.2,
        messages=[{"role": "user", "content": prompt}]
    )

    text = resp.choices[0].message.content or ""
    parsed = parse_ai(text)

    # zorunlu uyarı
    for k in parsed:
        parsed[k] += "\n\nBu bir hukuki tavsiye değildir."

    return parsed


@app.post("/export")
def export(req: ExportRequest, request: Request):
    tpl = get_template(req.template_key)
    if not tpl:
        raise HTTPException(400, "Şablon yok")

    prompt = build_prompt(tpl, req)

    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        temperature=0.2,
        messages=[{"role":"user","content":prompt}]
    )

    text = resp.choices[0].message.content or ""
    parsed = parse_ai(text)

    # ----------- DOCX ----------- #
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    p = doc.add_paragraph(tpl["title"])
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.bold = True

    def add_block(title, content):
        ph = doc.add_paragraph(title)
        ph.runs[0].bold = True
        for part in content.split("\n"):
            doc.add_paragraph(part.strip())

    add_block("Başlık", parsed["header"])
    add_block("Özet", parsed["summary"])
    add_block("Hukuki Sebepler", parsed["legal_basis"])
    add_block("Sonuç ve İstem", parsed["result_requests"])
    add_block("Ekler", parsed["attachments"])

    # otomatik eklenen zorunlu uyarı
    doc.add_paragraph("\nBu bir hukuki tavsiye değildir.")

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)

    fname = f"avo_{req.template_key}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"

    return StreamingResponse(
        out,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={fname}"}
    )
